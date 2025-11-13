import os
import json
import boto3
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

# AWS Clients
dynamodb = boto3.resource('dynamodb')
s3 = boto3.client('s3')

# Configuration
MEDICAL_HISTORIES_TABLE = os.environ.get('DYNAMODB_MEDICAL_HISTORIES_TABLE', 'medical_histories')
S3_BUCKET = os.environ.get('S3_EXPORTS_BUCKET', 'clinicalops-exports')
PRESIGNED_URL_EXPIRATION = 3600  # 1 hour

medical_histories_table = dynamodb.Table(MEDICAL_HISTORIES_TABLE)


def lambda_handler(event, context):
    """
    Export medical record to PDF or DOCX format.

    Expected payload:
    {
        "pathParameters": {
            "historyID": "string"
        },
        "queryStringParameters": {
            "format": "pdf" | "docx"
        }
    }
    """
    try:
        # Parse parameters
        path_params = event.get('pathParameters', {})
        query_params = event.get('queryStringParameters', {}) or {}

        history_id = path_params.get('historyID')
        export_format = query_params.get('format', 'pdf').lower()

        if not history_id:
            return {
                'statusCode': 400,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'historyID is required'})
            }

        if export_format not in ['pdf', 'docx']:
            return {
                'statusCode': 400,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'format must be "pdf" or "docx"'})
            }

        # Get medical record
        response = medical_histories_table.get_item(Key={'historyID': history_id})

        if 'Item' not in response:
            return {
                'statusCode': 404,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'Medical history not found'})
            }

        record = response['Item']
        structured_note = record.get('structuredClinicalNote', '')

        if not structured_note:
            return {
                'statusCode': 400,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'Medical record has no clinical note'})
            }

        # Parse JSON
        try:
            clinical_data = json.loads(structured_note)
        except json.JSONDecodeError:
            return {
                'statusCode': 500,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'Invalid JSON in clinical note'})
            }

        # Generate file
        if export_format == 'pdf':
            file_path = generate_pdf(history_id, record, clinical_data)
            content_type = 'application/pdf'
        else:  # docx
            file_path = generate_docx(history_id, record, clinical_data)
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        # Upload to S3
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        s3_key = f"exports/{history_id}/{timestamp}.{export_format}"

        with open(file_path, 'rb') as file_data:
            s3.put_object(
                Bucket=S3_BUCKET,
                Key=s3_key,
                Body=file_data,
                ContentType=content_type
            )

        print(f"Uploaded export to s3://{S3_BUCKET}/{s3_key}")

        # Generate presigned URL
        download_url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET, 'Key': s3_key},
            ExpiresIn=PRESIGNED_URL_EXPIRATION
        )

        # Clean up temp file
        os.remove(file_path)

        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({
                'message': 'Export generated successfully',
                'historyID': history_id,
                'format': export_format,
                'downloadUrl': download_url,
                'expiresIn': PRESIGNED_URL_EXPIRATION
            })
        }

    except Exception as e:
        print(f"Error exporting medical record: {e}")
        import traceback
        traceback.print_exc()

        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({'error': f'Internal server error: {str(e)}'})
        }


def generate_pdf(history_id, record, clinical_data):
    """Generate PDF document from clinical data"""
    # Create temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.close()

    # Create PDF
    doc = SimpleDocTemplate(temp_file.name, pagesize=letter,
                            rightMargin=0.75*inch, leftMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=12,
        alignment=1  # Center
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#333333'),
        spaceAfter=6,
        spaceBefore=12
    )
    body_style = styles['BodyText']

    # Build content
    story = []

    # Title
    story.append(Paragraph("Historia Clínica", title_style))
    story.append(Spacer(1, 0.2*inch))

    # Patient info
    patient_name = record.get('patientName', 'N/A')
    story.append(Paragraph(f"<b>Paciente:</b> {patient_name}", body_style))
    story.append(Paragraph(f"<b>ID:</b> {history_id}", body_style))
    story.append(Paragraph(f"<b>Fecha:</b> {datetime.now().strftime('%d/%m/%Y')}", body_style))
    story.append(Spacer(1, 0.3*inch))

    # Clinical data
    def add_section(key, value, level=0):
        """Recursively add sections to PDF"""
        indent = level * 20

        # Format key
        formatted_key = key.replace('_', ' ').title()

        if isinstance(value, dict):
            story.append(Paragraph(f"<b>{formatted_key}</b>", heading_style))
            for sub_key, sub_value in value.items():
                add_section(sub_key, sub_value, level + 1)
        elif isinstance(value, list):
            story.append(Paragraph(f"<b>{formatted_key}:</b>", heading_style))
            for item in value:
                if isinstance(item, dict):
                    for sub_key, sub_value in item.items():
                        add_section(sub_key, sub_value, level + 1)
                else:
                    story.append(Paragraph(f"• {item}", body_style))
        else:
            story.append(Paragraph(f"<b>{formatted_key}:</b> {value}", body_style))

    for key, value in clinical_data.items():
        add_section(key, value)

    # Build PDF
    doc.build(story)

    return temp_file.name


def generate_docx(history_id, record, clinical_data):
    """Generate DOCX document from clinical data"""
    # Create temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_file.close()

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading('Historia Clínica', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Patient info
    patient_name = record.get('patientName', 'N/A')
    doc.add_paragraph(f"Paciente: {patient_name}", style='Normal')
    doc.add_paragraph(f"ID: {history_id}", style='Normal')
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}", style='Normal')
    doc.add_paragraph('')  # Space

    # Clinical data
    def add_section(key, value, level=1):
        """Recursively add sections to DOCX"""
        # Format key
        formatted_key = key.replace('_', ' ').title()

        if isinstance(value, dict):
            doc.add_heading(formatted_key, level=level)
            for sub_key, sub_value in value.items():
                add_section(sub_key, sub_value, level + 1)
        elif isinstance(value, list):
            doc.add_heading(formatted_key, level=level)
            for item in value:
                if isinstance(item, dict):
                    for sub_key, sub_value in item.items():
                        add_section(sub_key, sub_value, level + 1)
                else:
                    doc.add_paragraph(f"• {item}", style='List Bullet')
        else:
            p = doc.add_paragraph()
            p.add_run(f"{formatted_key}: ").bold = True
            p.add_run(str(value))

    for key, value in clinical_data.items():
        add_section(key, value)

    # Save
    doc.save(temp_file.name)

    return temp_file.name
